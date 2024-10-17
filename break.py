

# import re
# from docx import Document
# from docx.shared import Inches

# # Load the document
# doc = Document("template.docx")

# # Create a new document to store the processed content
# new_doc = Document()

# # Add the image to the header
# section = new_doc.sections[0]
# header = section.header
# header_paragraph = header.paragraphs[0]
# header_run = header_paragraph.add_run()
# header_run.add_picture("logos_proj.jpeg", width=Inches(1.5))  # Adjust the path and size as needed

# # Function to extract headings using regex
# def extract_headings(paragraphs):
#     headings = []
#     for para in paragraphs:
#         if re.match(r'^\*\*.*\*\*$', para.text.strip()):  # Check if the paragraph is a heading
#             headings.append(para.text.strip().strip('*'))
#     return headings

# # Extract headings and subheadings from the document
# all_headings = extract_headings(doc.paragraphs)
# print("all_headings:", all_headings)

# # Ensure there are enough headings to prevent IndexError
# if len(all_headings) >= 2:
#     first_heading = all_headings[0]
#     second_heading = all_headings[1]
#     headings = all_headings[2:]
#     print("first_heading:", first_heading)
#     print("second_heading:", second_heading)
#     print("headings:", headings)
# else:
#     print("Not enough headings found in the document.")
#     first_heading = ""
#     second_heading = ""
#     headings = []

# # Add the custom headings at the top of the first page
# if first_heading:
#     new_doc.add_heading(first_heading, level=1)
# if second_heading:
#     new_doc.add_heading(second_heading, level=1)

# # Initialize variables
# current_heading = None
# is_collecting = False

# # Iterate over each paragraph in the document
# for para in doc.paragraphs:
#     para_text = para.text.strip()
#     print("para_text:", para_text)
#     # Check if the paragraph text matches any of the predefined headings
#     if para_text.strip('*') in headings:
#         # If we encounter a new heading, store it and start collecting content
#         if current_heading:
#             # Add a page break before moving to the next section
#             new_doc.add_page_break()
#         current_heading = para_text.strip('*')
#         is_collecting = True
#         # Add the heading to the new document
#         new_doc.add_heading(current_heading, level=1)
#         print("current_heading:", current_heading)
#     elif is_collecting:
#         # If we are in a section (after a heading), add the content
#         new_doc.add_paragraph(para.text)
#         print("text:", para.text)

# # Save the new document with dynamically identified sections
# new_doc.save(f"{first_heading}.docx")

# print("Document created successfully.")

# ## ***************** Approach 2

# import re
# from docx import Document
# from docx.shared import Inches

# # Load the document
# doc = Document("template.docx")

# # Create a new document to store the processed content
# new_doc = Document()

# # Add the image to the header
# section = new_doc.sections[0]
# header = section.header
# header_paragraph = header.paragraphs[0]
# header_run = header_paragraph.add_run()
# header_run.add_picture("logos_proj.jpeg", width=Inches(1.5))  # Adjust the path and size as needed

# # Function to extract headings using regex
# def extract_headings(paragraphs):
#     headings = []
#     for para in paragraphs:
#         if re.match(r'^\*\*.*\*\*$', para.text.strip()):  # Check if the paragraph is a heading
#             headings.append(para.text.strip().strip('*'))
#     return headings

# # Extract headings and subheadings from the document
# all_headings = extract_headings(doc.paragraphs)
# print("all_headings:", all_headings)

# # Ensure there are enough headings to prevent IndexError
# if len(all_headings) >= 2:
#     first_heading = all_headings[0]
#     second_heading = all_headings[1]
#     headings = all_headings[2:]
#     print("first_heading:", first_heading)
#     print("second_heading:", second_heading)
#     print("headings:", headings)
# else:
#     print("Not enough headings found in the document.")
#     first_heading = ""
#     second_heading = ""
#     headings = []

# # Add the custom headings at the top of the first page
# if first_heading:
#     new_doc.add_heading(first_heading, level=1)
# if second_heading:
#     new_doc.add_heading(second_heading, level=1)

# # Function to check if a heading includes a month (e.g., "For September")
# def is_month_heading(text):
#     return re.search(r'For\s\w+:', text)

# # Initialize variables
# current_heading = None
# is_collecting = False
# collecting_for_month = False  # Track if we are in a "For <Month>" section

# # Iterate over each paragraph in the document
# for para in doc.paragraphs:
#     para_text = para.text.strip()
#     print("para_text:", para_text)
    
#     # Check if the paragraph is a month heading (e.g., "For September:")
#     if is_month_heading(para_text):
#         # If we encounter a new month heading, we do NOT want to add a page break
#         current_heading = para_text
#         is_collecting = True
#         collecting_for_month = True
#         # Add the month heading to the new document
#         new_doc.add_heading(current_heading, level=1)
#         print("current_heading (month):", current_heading)
        
#     # Check if the paragraph text matches any of the predefined subheadings
#     elif para_text.strip('*') in headings:
#         # If we encounter a new subheading and we're not in a month section, add a page break
#         if current_heading and not collecting_for_month:
#             new_doc.add_page_break()
#         current_heading = para_text.strip('*')
#         is_collecting = True
#         collecting_for_month = False  # End of month-specific collection
#         # Add the subheading to the new document
#         new_doc.add_heading(current_heading, level=1)
#         print("current_heading:", current_heading)
        
#     elif is_collecting:
#         # If we are in a section (after a heading), add the content
#         new_doc.add_paragraph(para.text)
#         print("text:", para.text)

# # Save the new document with dynamically identified sections
# new_doc.save(f"{first_heading}.docx")

# print("Document created successfully.")

import re
from docx import Document
from docx.shared import Inches

# Load the document
doc = Document("colby break.docx")

# Create a new document to store the processed content
new_doc = Document()

# Add the image to the header
section = new_doc.sections[0]
header = section.header
header_paragraph = header.paragraphs[0]
header_run = header_paragraph.add_run()
header_run.add_picture("logos_proj.jpeg", width=Inches(1.5))  # Adjust the path and size as needed

# Function to extract headings using regex
def extract_headings(paragraphs):
    headings = []
    for para in paragraphs:
        if re.match(r'^\*\*.*\*\*$', para.text.strip()):  # Check if the paragraph is a heading
            headings.append(para.text.strip().strip('*'))
    return headings

# Extract headings and subheadings from the document
all_headings = extract_headings(doc.paragraphs)
print("all_headings:", all_headings)

# Ensure there are enough headings to prevent IndexError
if len(all_headings) >= 3:
    first_heading = all_headings[0]
    second_heading = all_headings[1]
    third_heading = all_headings[2]
    headings = all_headings[3:]
    print("first_heading:", first_heading)
    print("second_heading:", second_heading)
    print("third_heading:", third_heading)
    print("headings:", headings)
else:
    print("Not enough headings found in the document.")
    first_heading = ""
    second_heading = ""
    third_heading = ""
    headings = []

# Add the custom headings at the top of the first page
if first_heading:
    new_doc.add_heading(first_heading, level=1)
if second_heading:
    new_doc.add_heading(second_heading, level=1)
if third_heading:
    new_doc.add_heading(third_heading, level=1)

collecting_content = True  # Start collecting content after initial headings

# Initialize variables
current_heading = None

# Function to check if a heading is a month heading
def is_month_heading(text):
    return re.match(r'^For\s\w+:', text)

# Function to check if the text is a heading
def is_heading(text):
    return re.match(r'^\*\*.*\*\*$', text.strip())

# Function to get the heading text without asterisks
def get_heading_text(text):
    return text.strip().strip('*')

# Function to check if heading is 'Talking Points'
def is_talking_points(heading_text):
    return heading_text == 'Talking Points'

# Function to check if heading is 'Social Media Topic Ideas'
def is_social_media_topic_ideas(heading_text):
    return heading_text == 'Social Media Topic Ideas'

# Function to check if heading is 'Text Messaging Talking Points'
def is_text_messaging_talking_points(heading_text):
    return heading_text == 'Text Messaging Talking Points'

# Iterate over each paragraph in the document
for para in doc.paragraphs:
    para_text = para.text.strip()
    print("para_text:", para_text)

    # Check if the paragraph is a heading
    if is_heading(para_text):
        heading_text = get_heading_text(para_text)
        print("heading_text:", heading_text)

        if heading_text in [first_heading, second_heading, third_heading]:
            # These headings are already added on the first page
            # So we skip adding them again but start collecting content if it's 'TRS Messages'
            collecting_content = True
            current_heading = heading_text
            print("Skipped heading already added:", heading_text)
            continue

        # If it's a month heading
        elif is_month_heading(heading_text):
            # Always add a page break before month headings
            new_doc.add_page_break()
            # Add the month heading
            new_doc.add_heading(heading_text, level=1)
            current_heading = heading_text
            collecting_content = True
            print("Added month heading:", heading_text)

        elif is_talking_points(heading_text):
            # Add 'Talking Points' heading without page break
            new_doc.add_heading(heading_text, level=1)
            current_heading = heading_text
            collecting_content = True
            print("Added 'Talking Points' heading:", heading_text)

        elif is_social_media_topic_ideas(heading_text) or is_text_messaging_talking_points(heading_text):
            # Add a page break before these headings
            new_doc.add_page_break()
            new_doc.add_heading(heading_text, level=1)
            current_heading = heading_text
            collecting_content = True
            print("Added heading with page break:", heading_text)

        else:
            # For any other heading, add a page break and then the heading
            new_doc.add_page_break()
            new_doc.add_heading(heading_text, level=1)
            current_heading = heading_text
            collecting_content = True
            print("Added other heading with page break:", heading_text)

    else:
        # It's a normal paragraph
        if collecting_content:
            new_doc.add_paragraph(para.text)
            print("Added paragraph:", para.text)
        else:
            # Skip content before any headings
            print("Skipped paragraph before any heading")

# Save the new document with dynamically identified sections
new_doc.save(f"lmnop.docx")

print("Document created successfully.")






