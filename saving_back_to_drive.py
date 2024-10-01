

import streamlit as st
import docx
from docx import Document
import io
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from Google import Create_Service
import openai
import os
import re
import config
import zipfile
from  groq import Groq
import anthropic

# Set OpenAI key
openai.api_key = config.OPENAI_API_KEY
os.environ['GROQ_API_KEY']=config.GROQ_API_KEY
os.environ['ANTHROPIC_API_KEY']=config.ANTHROPIC_API_KEY


# Google Drive API setup
CLIENT_SECRET_FILE = 'client_secret.json'
API_NAME = 'drive'
API_VERSION = 'v3'
SCOPES = ['https://www.googleapis.com/auth/drive']
service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)

# Function to list files in folder
def list_files_in_folder(folder_id):
    results = service.files().list(
        q=f"'{folder_id}' in parents",
        spaces='drive',
        fields='files(id, name, mimeType)').execute()
    return results.get('files', [])

# Function to read .docx files
def read_docx(file_stream):
    document = Document(file_stream)
    content = []
    for paragraph in document.paragraphs:
        content.append(paragraph.text)
    return '\n'.join(content)

# Function to stream file content from Google Drive
def stream_file(service, file_id):
    request = service.files().get_media(fileId=file_id)
    file_stream = io.BytesIO()
    downloader = MediaIoBaseDownload(file_stream, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    file_stream.seek(0)
    return file_stream

# Function to create a folder in Google Drive
def create_drive_folder(service, folder_name, parent_folder_id=None):
    file_metadata = {
        'name': folder_name,
        'mimeType': 'application/vnd.google-apps.folder'
    }
    if parent_folder_id:
        file_metadata['parents'] = [parent_folder_id]
    
    folder = service.files().create(body=file_metadata, fields='id').execute()
    return folder['id']

# Function to upload a file to a specific folder in Google Drive
def upload_to_drive(service, file_name, file_path, folder_id):
    file_metadata = {
        'name': file_name,
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file['id']

# Function to generate prompt for GPT
def prompt_generator(content):
    prompt = f"""
      Using the information and instructions provided in {docx_content}, generate a message in the following format, including the headings (Talking Points, Social Media Topic Ideas, Text Messaging Talking Points).Ensure that each heading contains at least 5-6 bullet points.

      **<College> <Sport>**
      **Sept./Oct./Nov./Dec. 2024**
      **TRS Messages**

      For September, focus on residence halls and general everyday life on campus for both students and athletes. According to our research with your team at <college1> and other colleges across the country, this is a key area of interest for this generation of recruits.
      October: Highlight the overall athletic climate at <College>, showcasing what it’s like to compete and be part of the campus community as both an athlete and a student.
      November: Emphasize the athletic facilities and training philosophy at <College>, demonstrating how recruits will be prepared for college-level competition.
      December: Focus on the <sport> team atmosphere at <College>, incorporating insights from the focus group survey to explore team dynamics.

     Make sure that each month include the headings (Talking Points, Social Media Topic Ideas, Text Messaging Talking Points).
"""

    return prompt

def prompt_generator_for_sonnet(content):
    prompt_temp = f"""
Based on the provided content from {docx_content}, generate a recruiting message in the format below. Maintain the structure and include the specified headings, subheadings, and bullet points.

**<College> <Sport>**
**Sept./Oct./Nov./Dec. 2024**
**TRS Messages**

**For September:**
Highlight residence halls and daily student-athlete life. This aligns with research from <College1> and other institutions.

**Talking Points (5-6 bullet points)**
**Social Media Topic Ideas (5-6 bullet points)**
**Text Messaging Talking Points (5-6 bullet points)**

**For October:**
Showcase the athletic experience at <College>—what it’s like being part of the athletic community.

**Talking Points (5-6 bullet points)**
**Social Media Topic Ideas (5-6 bullet points)**
**Text Messaging Talking Points (5-6 bullet points)**

**For November:**
Emphasize facilities and training that prepare recruits for college-level competition.

**Talking Points (5-6 bullet points)**
**Social Media Topic Ideas (5-6 bullet points)**
**Text Messaging Talking Points (5-6 bullet points**

**For December:**
Focus on team dynamics and atmosphere. Include relevant insights from focus group surveys.

**Talking Points (5-6 bullet points)**
**Social Media Topic Ideas (5-6 bullet points)**
**Text Messaging Talking Points (5-6 bullet points)**

Ensure that the headings, subheadings, and bullet points remain organized in the final output.
"""
    return prompt_temp

# Function to format and save content in .docx
def format_content(doc, content):
    lines = content.splitlines()
    for line in lines:
        line = line.strip()
        if not line:
            continue
        if re.match(r'^\*\*.+\*\*$', line):
            doc.add_heading(line.replace('**', ''), level=1)
        elif re.match(r'^\*.+\*$', line):
            doc.add_heading(line.replace('*', ''), level=2)
        elif line.startswith("-"):
            doc.add_paragraph(line, style='List Bullet')
        else:
            doc.add_paragraph(line)



def zip_folder(folder_path, output_path):
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                # Preserve folder structure inside the zip
                arcname = os.path.relpath(file_path, start=folder_path)
                zipf.write(file_path, arcname=arcname)

# Streamlit UI
st.title("Content Generator")
group = st.text_input("Enter group name:")
model_selection = st.selectbox("Choose Model", ["ChatGPT-4", "LLaMA3", "Sonnet"])
run_process = st.button("Generate Responses")

if run_process and group:
    # List files in the group folder
    response = service.files().list(q=f"name = '{group}' and mimeType = 'application/vnd.google-apps.folder'",
                                    spaces='drive').execute()
    folders = response.get('files', [])

    if not folders:
        st.error("Folder not found.")
    else:
        folder_id = folders[0]['id']
        response_folder = os.path.join(f'{group}_responses', group)  # Create Group A folder inside responses
        if not os.path.exists(response_folder):
            os.makedirs(response_folder)

        # Create a new folder 'Group A responses' in Google Drive
        responses_folder_id = create_drive_folder(service, 'Group A responses')

       
        subfolders = list_files_in_folder(folder_id)
        for folder in subfolders:
            if folder['mimeType'] == 'application/vnd.google-apps.folder':
                st.write(f"Subfolder: {folder['name']}")
                files_in_subfolder = list_files_in_folder(folder['id'])

                for file in files_in_subfolder:
                    if file['name'].endswith('.docx'):
                        st.write(f"Processing File: {file['name']}")
                        file_stream = stream_file(service, file['id'])
                        docx_content = read_docx(file_stream)

                        # Generate GPT prompt and response
                        prompt_template = prompt_generator(docx_content)
                        promptt = prompt_generator_for_sonnet(docx_content)
                        if model_selection == "ChatGPT-4":
                            response = openai.ChatCompletion.create(
                                model="gpt-4-turbo",
                                messages=[
                                    {"role": "system", "content": "You are a helpful assistant."},
                                    {"role": "user", "content": prompt_template}
                                ],
                                max_tokens=2000,
                                temperature=0.5
                            )
                            result = response.choices[0].message['content']
                            print(result)
                            # Save response in .docx
                            college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                            doc = Document()
                            format_content(doc, result)
                            doc.save(college_response_file)
                            st.write(f"Saved: {college_response_file}")
                            uploaded_file_id = upload_to_drive(service, f"{folder['name']} response.docx", college_response_file, responses_folder_id)
                            st.write(f"Uploaded: {folder['name']} response.docx")

                        elif model_selection == "LLaMA3":
                            client = Groq(api_key=os.getenv("GROQ_API_KEY"))
                            chat_completion = client.chat.completions.create(
                                messages=[
                                    {
                                        "role": "user",
                                        "content": prompt_template,
                                    }
                                ],
                                model="llama3-8b-8192",
                            )
                            result = chat_completion.choices[0].message.content
                            print(result)
                            # Save response in .docx
                            college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                            doc = Document()
                            format_content(doc, result)
                            doc.save(college_response_file)
                            st.write(f"Saved: {college_response_file}")
                            uploaded_file_id = upload_to_drive(service, f"{folder['name']} response.docx", college_response_file, responses_folder_id)
                            st.write(f"Uploaded: {folder['name']} response.docx")

                        elif model_selection == "Sonnet":
                            client = anthropic.Anthropic(
                                api_key=os.getenv("ANTHROPIC_API_KEY"),
                            )
                            message = client.messages.create(
                                model="claude-3-5-sonnet-20240620",
                                max_tokens=2000,
                                temperature=0,
                                system="You are a helpful assistant for generating long text.",
                                messages=[
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": promptt
                                            }
                                        ]
                                    }
                                ]
                            )
                            result = message.content[0].text
                            print(result)
                            # Save response in .docx
                            college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                            doc = Document()
                            format_content(doc, result)
                            doc.save(college_response_file)
                            st.write(f"Saved: {college_response_file}")
                            # Upload the generated .docx file to the 'Group A responses' folder in Google Drive
                            uploaded_file_id = upload_to_drive(service, f"{folder['name']} response.docx", college_response_file, responses_folder_id)
                            st.write(f"Uploaded: {folder['name']} response.docx")

                        # # Save response in .docx inside the Group A folder
                        # college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                        # doc = Document()
                        # format_content(doc, result)
                        # doc.save(college_response_file)
                        # st.write(f"Saved: {college_response_file}")

        # Zip the entire Group A folder
        zip_file_path = f"{group}_responses.zip"
        zip_folder(os.path.join(f'{group}_responses'), zip_file_path)

        # Provide download button for the ZIP file
        with open(zip_file_path, "rb") as f:
            st.download_button(
                label=f"Download {group} responses (ZIP)",
                data=f,
                file_name=f"{group}_responses.zip",
                mime="application/zip"
            )




