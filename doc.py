# import streamlit as st
# import docx
# # from docx import Document
# from googleapiclient.http import MediaIoBaseDownload
# from Google import Create_Service
# import openai
# import io
# import os
# from decouple import config

# # Configure your OpenAI API Key
# # os.environ['OPENAI_API_KEY'] = config.OPENAI_API_KEY
# openai.api_key = config('OPENAI_API_KEY')

# # Google Drive API Configuration
# CLIENT_SECRET_FILE = 'client_secret.json'  # Path to your client secret file
# API_NAME = 'drive'
# API_VERSION = 'v3'
# SCOPES = ['https://www.googleapis.com/auth/drive']
# service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)

# # Streamlit App Layout
# st.title("Content Generator")
# group = st.text_input("Enter group name (e.g., Group A):")

# # Function to list files in a folder
# def list_files_in_folder(folder_id):
#     results = service.files().list(
#         q=f"'{folder_id}' in parents",
#         spaces='drive',
#         fields='files(id, name, mimeType)').execute()
#     return results.get('files', [])

# # Function to read .docx content
# def read_docx(file_stream):
#     document = docx.Document(file_stream)
#     content = []
#     for paragraph in document.paragraphs:
#         content.append(paragraph.text)
#     return '\n'.join(content)

# # Function to stream a file from Google Drive
# def stream_file(service, file_id):
#     request = service.files().get_media(fileId=file_id)
#     file_stream = io.BytesIO()
#     downloader = MediaIoBaseDownload(file_stream, request)
#     done = False
#     while not done:
#         status, done = downloader.next_chunk()
#         # print(f"Download {int(status.progress() * 100)}%.")
#     file_stream.seek(0)
#     return file_stream

# # Function to generate the prompt for OpenAI
# def prompt_generator(content):
#     prompt = f"""
#     Using the information and instructions provided in {content}, generate a message in the following format, including the headings (Talking Points, Social Media Topic Ideas, Text Messaging Talking Points). Ensure that each heading contains at least 5-6 bullet points.

#     <College> <Sport>
#     Sept./Oct./Nov./Dec. 2024
#     TRS Messages

#     For September, focus on residence halls and general everyday life on campus for both students and athletes. According to our research with your team at <college1> and other colleges across the country, this is a key area of interest for this generation of recruits.
#     October: Highlight the overall athletic climate at <College>, showcasing what it’s like to compete and be part of the campus community as both an athlete and a student.
#     November: Emphasize the athletic facilities and training philosophy at <College>, demonstrating how recruits will be prepared for college-level competition.
#     December: Focus on the <sport> team atmosphere at <College>, incorporating insights from the focus group survey to explore team dynamics.

#     Make sure that each month includes the headings (Talking Points, Social Media Topic Ideas, Text Messaging Talking Points).
#     """
#     return prompt



# if st.button("Fetch and Generate"):
#     if not group:
#         st.error("Please enter a group name.")
#     else:
#         response = service.files().list(q=f"name = '{group}' and mimeType = 'application/vnd.google-apps.folder'",
#                                         spaces='drive').execute()
#         folders = response.get('files', [])

#         if not folders:
#             st.error("Folder not found.")
#         else:
#             folder_id = folders[0]['id']
#             subfolders = list_files_in_folder(folder_id)
#             # st.write(f"Folder: {group}")

#             for folder in subfolders:
#                 if folder['mimeType'] == 'application/vnd.google-apps.folder':
#                     # st.write(f"Subfolder: {folder['name']}")
#                     files_in_subfolder = list_files_in_folder(folder['id'])
#                     print(files_in_subfolder)
                    
#                     for file in files_in_subfolder:
#                         if file['name'].endswith('.docx'):
#                             st.write(f"Processing File: {file['name']}")
#                             file_stream = stream_file(service, file['id'])
#                             docx_content = read_docx(file_stream)

#                             # Generate prompt and send to OpenAI
#                             prompt_template = prompt_generator(docx_content)
#                             response = openai.ChatCompletion.create(
#                                 model="gpt-4-turbo",
#                                 messages=[
#                                     {"role": "system", "content": "You are a helpful assistant."},
#                                     {"role": "user", "content": prompt_template}
#                                 ],
#                                 max_tokens=2000,
#                                 temperature=0.5
#                             )

#                             # Display OpenAI response
#                             st.write(f"Response for {file['name']}:")
#                             st.write(response.choices[0].message['content'])
#                             st.write("=" * 80)

import streamlit as st
import docx
from googleapiclient.http import MediaIoBaseDownload
from Google import Create_Service
import openai
import io
import os
from decouple import config
from PyPDF2 import PdfReader
from langchain.text_splitter import CharacterTextSplitter
from langchain.embeddings.openai import OpenAIEmbeddings
from langchain.vectorstores import Chroma

# Configure your OpenAI API Key
openai.api_key = config('OPENAI_API_KEY')

# Google Drive API Configuration
CLIENT_SECRET_FILE = 'client_secret.json'  # Path to your client secret file
API_NAME = 'drive'
API_VERSION = 'v3'
SCOPES = ['https://www.googleapis.com/auth/drive']
service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)

# Streamlit App Layout
st.title("Content Generator")
group = st.text_input("Enter group name (e.g., Group A):")

# Function to list files in a folder
def list_files_in_folder(folder_id):
    results = service.files().list(
        q=f"'{folder_id}' in parents",
        spaces='drive',
        fields='files(id, name, mimeType)').execute()
    return results.get('files', [])

# Function to read .docx content
def read_docx(file_stream):
    document = docx.Document(file_stream)
    content = []
    for paragraph in document.paragraphs:
        content.append(paragraph.text)
    return '\n'.join(content)

# Function to read .pdf content
def read_pdf(file_stream):
    pdf_reader = PdfReader(file_stream)
    raw_text = ''
    for page in pdf_reader.pages:
        text = page.extract_text()
        if text:
            raw_text += text
    return raw_text

# Function to stream a file from Google Drive
def stream_file(service, file_id):
    request = service.files().get_media(fileId=file_id)
    file_stream = io.BytesIO()
    downloader = MediaIoBaseDownload(file_stream, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    file_stream.seek(0)
    return file_stream

# Function to generate the prompt for OpenAI
def prompt_generator(content):
    prompt = f"""
    Using the information and instructions provided in {content}, generate a message in the following format, including the headings (Talking Points, Social Media Topic Ideas, Text Messaging Talking Points). Ensure that each heading contains at least 5-6 bullet points.

    <College> <Sport>
    Sept./Oct./Nov./Dec. 2024
    TRS Messages

    For September, focus on residence halls and general everyday life on campus for both students and athletes. According to our research with your team at <college1> and other colleges across the country, this is a key area of interest for this generation of recruits.
    October: Highlight the overall athletic climate at <College>, showcasing what it’s like to compete and be part of the campus community as both an athlete and a student.
    November: Emphasize the athletic facilities and training philosophy at <College>, demonstrating how recruits will be prepared for college-level competition.
    December: Focus on the <sport> team atmosphere at <College>, incorporating insights from the focus group survey to explore team dynamics.

    Make sure that each month includes the headings (Talking Points, Social Media Topic Ideas, Text Messaging Talking Points).
    """
    return prompt

# Function to split text and store in Chroma DB
def store_in_chroma(text_content, file_name):
    # Clean the file name to use as the directory name
    persist_directory = file_name.replace(' ', '_').replace('.pdf', '').replace('.docx', '')

    text_splitter = CharacterTextSplitter(
        separator = '\n',
        chunk_size = 800,
        chunk_overlap = 200,
        length_function = len,
    )
    
    # Split the text into manageable chunks
    texts = text_splitter.split_text(text_content)
    
    # Generate embeddings using OpenAI
    embeddings = OpenAIEmbeddings()
    
    # Create a Chroma DB and persist the data in a directory named after the file
    chroma_db = Chroma.from_texts(texts, embeddings, persist_directory=persist_directory)
    chroma_db.persist()

# Main processing logic
def process_files(files_in_subfolder):
    for file in files_in_subfolder:
        if file['name'].endswith(('.docx', '.pdf')):
            st.write(f"Processing File: {file['name']}")
            file_stream = stream_file(service, file['id'])

            if file['name'].endswith('.docx'):
                docx_content = read_docx(file_stream)
                content = docx_content  # Use the content extracted from .docx

            elif file['name'].endswith('.pdf'):
                pdf_content = read_pdf(file_stream)
                content = pdf_content  # Use the content extracted from .pdf

            # Store extracted content in Chroma DB with file-based directory
            store_in_chroma(content, file['name'])

            # Generate prompt and send to OpenAI
            prompt_template = prompt_generator(content)
            response = openai.ChatCompletion.create(
                model="gpt-4-turbo",
                messages=[
                    {"role": "system", "content": "You are a helpful assistant."},
                    {"role": "user", "content": prompt_template}
                ],
                max_tokens=2000,
                temperature=0.5
            )

            # Display OpenAI response
            st.write(f"Response for {file['name']}:")
            st.write(response.choices[0].message['content'])
            st.write("=" * 80)

# Main Streamlit logic to fetch files and process them
if st.button("Fetch and Generate"):
    if not group:
        st.error("Please enter a group name.")
    else:
        response = service.files().list(q=f"name = '{group}' and mimeType = 'application/vnd.google-apps.folder'",
                                        spaces='drive').execute()
        folders = response.get('files', [])

        if not folders:
            st.error("Folder not found.")
        else:
            folder_id = folders[0]['id']
            subfolders = list_files_in_folder(folder_id)

            for folder in subfolders:
                if folder['mimeType'] == 'application/vnd.google-apps.folder':
                    files_in_subfolder = list_files_in_folder(folder['id'])
                    process_files(files_in_subfolder)
