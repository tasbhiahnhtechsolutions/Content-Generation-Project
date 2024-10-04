from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
document = Document()
header = document.sections[0].header
htable=header.add_table(1, 2, Inches(6))
htab_cells=htable.rows[0].cells
ht0=htab_cells[0].add_paragraph()
kh=ht0.add_run()
kh.add_picture('logos_proj.jpeg', width=Inches(1))
# ht1=htab_cells[1].add_paragraph('put your header text here')
# ht1.alignment = WD_ALIGN_PARAGRAPH.RIGHT
document.save('yourdoc.docx')