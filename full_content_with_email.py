

import streamlit as st
import docx
from docx import Document
import io
from googleapiclient.http import MediaIoBaseDownload, MediaFileUpload
from Google import Create_Service
import openai
import os
import re
from decouple import config
import zipfile
from  groq import Groq
import anthropic
from docx.shared import Inches, Pt, RGBColor


# Set OpenAI key
openai.api_key = config('OPENAI_API_KEY')
GROQ_API_KEY=config('GROQ_API_KEY')
ANTHROPIC_API_KEY=config('ANTHROPIC_API_KEY')


# Google Dr)ve API setup
CLIENT_SECRET_FILE = 'client_secret.json'
API_NAME = 'drive'
API_VERSION = 'v3'
SCOPES = ['https://www.googleapis.com/auth/drive']
service = Create_Service(CLIENT_SECRET_FILE, API_NAME, API_VERSION, SCOPES)


# Function to list files in folder
def list_files_in_folder(folder_id):
    results = service.files().list(
        q=f"'{folder_id}' in parents",
        spaces='drive',
        fields='files(id, name, mimeType)').execute()
    return results.get('files', [])

# Function to read .docx files
def read_docx(file_stream):
    document = Document(file_stream)
    content = []
    for paragraph in document.paragraphs:
        content.append(paragraph.text)
    return '\n'.join(content)

# Function to stream file content from Google Drive
def stream_file(service, file_id):
    request = service.files().get_media(fileId=file_id)
    file_stream = io.BytesIO()
    downloader = MediaIoBaseDownload(file_stream, request)
    done = False
    while not done:
        status, done = downloader.next_chunk()
    file_stream.seek(0)
    return file_stream

# Function to create a folder in Google Drive
def create_drive_folder(service, folder_name, parent_folder_id=None):
    file_metadata = {
        'name': folder_name,
        'mimeType': 'application/vnd.google-apps.folder'
    }
    if parent_folder_id:
        file_metadata['parents'] = [parent_folder_id]
    
    folder = service.files().create(body=file_metadata, fields='id').execute()
    return folder['id']

# Function to upload a file to a specific folder in Google Drive
def upload_to_drive(service, file_name, file_path, folder_id):
    file_metadata = {
        'name': file_name,
        'parents': [folder_id]
    }
    media = MediaFileUpload(file_path, mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    file = service.files().create(body=file_metadata, media_body=media, fields='id').execute()
    return file['id']




# Function to generate prompt for GPT
# def prompt_generator(content,user_input,full_months):
#     prompt = f"""
# Using the content from the provided document {docx_content}, generate a detailed recruiting message that follows this structure:

# 1. **College Name and Sport**: Clearly mention the college and sport at the top.
# 2. **Time Period "{user_input}"**: Include the time period of the message.
# 3. **TRS Messages**: Provide an overview of the monthly focus topics for each month. Each month should be randomly assigned one of the following topics:
# History and Vision for the Program
# Athletic Facilities
# Life After College
# Academics
# Athletic Atmosphere at the School
# Dorms and Campus Life
# Coaching
# The Freshman Experience
# Location and Area
# Our Team
# For each month, the TRS message should look like this format:

# In September The residence halls and general everyday life on campus for students and athletes will be the focus, based on your team’s feedback at <college1> and others nationwide. This is an important topic for this generation of recruits.
# In October: The athletic atmosphere at <college1> will be the focus, giving recruits an idea of what it’s like to compete and live as a student-athlete at <college1>.
# In November: The athletic facilities at <college1> will be highlighted, emphasizing how your training philosophy prepares athletes to compete at the collegiate level.
# In December: We’ll focus on the <sport> team at <college1>, including insights into the team atmosphere, based on recent findings from your focus group survey.
# *VERY IMPORTANT*: Ensure each section contains more content, with longer sentences under each heading to provide comprehensive information.

# For each month ({full_months[0]}, {full_months[1]}, {full_months[2]}, and {full_months[3]}), follow this structure:

# For [Month]:

# **Main Topic**: The randomly assigned topic for that month.
# **Talking Points (8 elaborated points in question form)**:  The talking points should explore the student's personal preferences, concerns, and expectations. Make sure the questions are open-ended and encourage the student to reflect on their experiences and aspirations. Use a friendly and informal tone, as if a college recruiter or coach is having a one-on-one conversation with the student.
# Make sure that all the points are detailed.
# **Social Media Topic Ideas (8 elaborated points)**: Generate creative social media post ideas and related social activities for engaging college students and athletes. For each social media platform suggestion, follow it up with a related social activity idea. Ensure that the points alternate between social media content and social activity. The platforms should include Instagram, Snapchat, Twitter (X), Facebook, YouTube, LinkedIn, Reddit, and TikTok. The goal is to balance digital engagement with in-person or virtual team-building activities.
# **Text Messaging Talking Points (8 elaborated points in question form)**: Create engaging questions recruiters can send to prospects via text message, tailored to the main topic of the month.
#  Draft an email or letter outline for each week of the month. Include:
# -Week 1: Introduction to the month’s theme and invitation for questions.
# -Week 2: Share personal stories or testimonials related to the theme.
# -Week 3: Provide valuable insights or resources (like facility tours).
# -Week 4: Encourage recruits to envision themselves in the program and invite further questions.
# Ensure each section contains longer sentences with detailed content that can be easily understood by teenagers. Make sure the headings, subheadings, and bullet points remain well-organized in the final output.

# **IMPORTANT** This is an Example Content for reference for Talking Points, Social Media Topic Ideas, and Text Messaging Talking Points:

# Talking Points Example:

# -What have your parents said when it comes to the idea of living on a college campus and being away at college?
# -Aside from <sport>, do you see yourself getting involved in any other aspects of college life? Have you thought about what you’ll be doing in between classes and practice?
# -What kind of atmosphere do you prefer when it comes to dorm life? (Are you a morning person or a night owl?)
# -What are your feelings about living away from home?
# -How do you picture college life?
# -Walk me through some of the things you’re a little nervous about:
# -Are you more of a private person? Shy or outgoing?
# -How do you feel about the idea of having a roommate?
# --Do you have any food allergies (or just preferences)? Are you a picky eater?

# Social Media Topic Ideas Example:

# -One picture a week inside the dorms - rooms, common areas, etc. Our studies show that your prospects need a peek at what they would see on campus as a way to get them to commit to visiting campus.
# -Encourage your team to get together and do a live stream on social media just for your recruits. Do it from where they live, and let them go around and tell recruits what it’s like on campus.
# -Include your team in as much as possible. Let them show off their dorm rooms and have some fun with it!
# -Try to get someone NOT associated with your team or athletic department to write a quick post with their picture, talking about life on campus and their role. Begin introducing your prospects to the people outside of their sport that they need to hear from.
# -Let your team know we’re talking about this and ask them to create some posts about it. Get them directly involved!
# -Twitter: Tweet your Top 20 short comments from your team about the dorm, spread out over the month. Have them tweet it, and then you retweet it.
# -Video focus: Where they’ll eat. Get a video of your team getting food, where they sit and eat, etc. Post that on all social media video outlets.
# -When it comes to the topic of where they live, your own athletes are the best at coming up with topics and visuals - rely on them to come up with ideas surrounding the space where they live.

# -Text Messaging Talking Points Example:

# -What do you think your life is going to look like when you aren’t practicing or in class? Have you thought about that yet?
# -[Coach: If there’s any particularly fun campus event that goes on, feel free to text your recruits about it. Include a picture!]
# -Reply back with your first instinct…now that you know a little bit about us, can you see yourself enjoying life on <college1>’s campus?
# -<Prospect Name>…what’s your favorite post-practice meal? My athletes love <mention food/snacks available on campus>.
# -What are some things you’d like to know about our dorms—the place where you would be living as a student-athlete here? [OR IF ALREADY VISITED: “I’m wondering…what do you remember about our dorms when you visited? Did you like them?”]
# Keep in mind that this is just for reference donot use its content in the content you will be generating.

# Ensure the output aligns with the template format below:

# **<College Name> <Sport>**
# **{user_input}**
# **TRS Messages**
# In {full_months[0]} Brief description about randomly assigned topic 1 in the same format as provide above.
# In {full_months[1]}: Brief description about randomly assigned topic 2 in the same format as provide above.
# In {full_months[2]}: Brief description about randomly assigned topic 3 in the same format as provide above.
# In {full_months[3]}: Brief description about randomly assigned topic 4 in the same format as provide above.
# **For {full_months[0]}: Main Topic**
# **Talking Points**
# **Social Media Topic Ideas**
# **Text Messaging Talking Points**
# The email should be written in a following format:
    
#     **{full_months[0]}/randomly assigned topic above**
#     **WEEK 1**
#     **Email 1**
#     **Suggested subject line**: Subject line
#     <Prospect Name>
#     Email content  
#     <Coach Info>
    



# **For {full_months[1]}: Main Topic**
# **Talking Points**
# **Social Media Topic Ideas**
# **Text Messaging Talking Points**
# The email should be written in a following format:
    
#     **{full_months[1]}/randomly assigned topic above**
#     **WEEK 1**
#     **Email 1**
#     **Suggested subject line**: Subject line
#     <Prospect Name>
#     Email content  
#     <Coach Info>

# **For {full_months[2]}: Main Topic**
# **Talking Points**
# **Social Media Topic Ideas**
# **Text Messaging Talking Points**
# The email should be written in a following format:
    
#     **{full_months[2]}/randomly assigned topic above**
#     **WEEK 1**
#     **Email 1**
#     **Suggested subject line**: Subject line
#     <Prospect Name>
#     Email content  
#     <Coach Info>

# **For {full_months[3]}: Main Topic**
# **Talking Points**
# **Social Media Topic Ideas**
# **Text Messaging Talking Points**
# The email should be written in a following format:
    
#     **{full_months[3]}/randomly assigned topic above**
#     **WEEK 1**
#     **Email 1**
#     **Suggested subject line**: Subject line
#     <Prospect Name>
#     Email content  
#     <Coach Info>

# Make the text more conversational, and write it in a way that would make a 16 to 18 year old teenager engage with the content and be prompted to respond and interact with the coach who is sending these
# Look for any grammatical errors and correct them.
# VERY IMPORTANT: Place the proper spacing between paragraphs in the revised text
# Keep the same layout format in place when you construct your revised text
# Use the content from the {docx_content} to fill in the placeholders for the talking points, social media ideas, and text messaging points.
# Ensure that the headings, subheadings, and bullet points remain organized in the final output.
# Make sure every college follows the same template structure.
# ONLY include the content necessary for generating the recruiting message. Do not add any extra or irrelevant details.




# """


def prompt_generator(content,user_input,full_months):
    prompt = f"""
Using the content from the provided document {docx_content} to fill in the placeholders only, generate a detailed recruiting message that follows this structure:

1. **College Name and Sport**: Clearly mention the college and sport at the top.
2. **Time Period "{user_input}"**: Include the time period of the message.
3. **TRS Messages**: Provide an overview of the monthly focus topics for each month. Each month should be randomly assigned one of the following topics:
History and Vision for the Program
Athletic Facilities
Life After College
Academics
Athletic Atmosphere at the School
Dorms and Campus Life
Coaching
The Freshman Experience
Location and Area
Our Team
For each month, the TRS message should look like this format:

In September The residence halls and general everyday life on campus for students and athletes will be the focus, based on your team’s feedback at <college1> and others nationwide. This is an important topic for this generation of recruits.
In October: The athletic atmosphere at <college1> will be the focus, giving recruits an idea of what it’s like to compete and live as a student-athlete at <college1>.
In November: The athletic facilities at <college1> will be highlighted, emphasizing how your training philosophy prepares athletes to compete at the collegiate level.
In December: We’ll focus on the <sport> team at <college1>, including insights into the team atmosphere, based on recent findings from your focus group survey.
*VERY IMPORTANT*: Ensure each section contains more content, with longer sentences under each heading to provide comprehensive information.

For each month ({full_months[0]}, {full_months[1]}, {full_months[2]}, and {full_months[3]}), follow this structure:

For [Month]:

**Main Topic**: The randomly assigned topic for that month.
**Talking Points (8 elaborated points in question form)**:  The talking points should explore the student's personal preferences, concerns, and expectations. Make sure the questions are open-ended and encourage the student to reflect on their experiences and aspirations. Use a friendly and informal tone, as if a college recruiter or coach is having a one-on-one conversation with the student.
Make sure that all the points are detailed.
**Social Media Topic Ideas (8 elaborated points)**: Generate creative social media post ideas and related social activities for engaging college students and athletes. For each social media platform suggestion, follow it up with a related social activity idea. Ensure that the points alternate between social media content and social activity. The platforms should include Instagram, Snapchat, Twitter (X), Facebook, YouTube, LinkedIn, Reddit, and TikTok. The goal is to balance digital engagement with in-person or virtual team-building activities.
**Text Messaging Talking Points (8 elaborated points in question form)**: Create engaging questions recruiters can send to prospects via text message, tailored to the main topic of the month.
Draft an email templates for the recruited candidate, one for Week 1, Week 2, and Week 4, along with a letter for the parent in Week 2 and a letter for the coach in Week 1. Each communication should align with the monthly outline:
Week 1 (Candidate Email 1): Introduce the month’s theme and invite the candidate to ask any questions they may have.
Week 2 (Candidate Email 2): Share personal stories or testimonials related to the theme to help the candidate feel more connected.
Week 2 (Parent Letter): Address concerns of the prospect's family, emphasizing the school's support system and the holistic development of the student-athlete.
Week 3 (Candidate Email 3): Encourage the candidate to envision themselves as part of the program, and invite any final questions they might have.
Week 4 (Coach Letter): A motivational letter from the coach to the student-athlete, providing insights into the program and expressing excitement for their potential involvement.
Week 4 (Letter 1): A direct, personal letter to the prospect, focusing on their aspirations and why the program fits their future.
Ensure each email and letter is detailed and written in a warm, conversational tone.

Ensure each section contains longer sentences with detailed content that can be easily understood by teenagers. Make sure the headings, subheadings, and bullet points remain well-organized in the final output.

**IMPORTANT** This is an Example Content for reference for Talking Points, Social Media Topic Ideas, and Text Messaging Talking Points:

Talking Points Example:

-What have your parents said when it comes to the idea of living on a college campus and being away at college?
-Aside from <sport>, do you see yourself getting involved in any other aspects of college life? Have you thought about what you’ll be doing in between classes and practice?
-What kind of atmosphere do you prefer when it comes to dorm life? (Are you a morning person or a night owl?)
-What are your feelings about living away from home?
-How do you picture college life?
-Walk me through some of the things you’re a little nervous about:
-Are you more of a private person? Shy or outgoing?
-How do you feel about the idea of having a roommate?
--Do you have any food allergies (or just preferences)? Are you a picky eater?

Social Media Topic Ideas Example:

-One picture a week inside the dorms - rooms, common areas, etc. Our studies show that your prospects need a peek at what they would see on campus as a way to get them to commit to visiting campus.
-Encourage your team to get together and do a live stream on social media just for your recruits. Do it from where they live, and let them go around and tell recruits what it’s like on campus.
-Include your team in as much as possible. Let them show off their dorm rooms and have some fun with it!
-Try to get someone NOT associated with your team or athletic department to write a quick post with their picture, talking about life on campus and their role. Begin introducing your prospects to the people outside of their sport that they need to hear from.
-Let your team know we’re talking about this and ask them to create some posts about it. Get them directly involved!
-Twitter: Tweet your Top 20 short comments from your team about the dorm, spread out over the month. Have them tweet it, and then you retweet it.
-Video focus: Where they’ll eat. Get a video of your team getting food, where they sit and eat, etc. Post that on all social media video outlets.
-When it comes to the topic of where they live, your own athletes are the best at coming up with topics and visuals - rely on them to come up with ideas surrounding the space where they live.

-Text Messaging Talking Points Example:

-What do you think your life is going to look like when you aren’t practicing or in class? Have you thought about that yet?
-[Coach: If there’s any particularly fun campus event that goes on, feel free to text your recruits about it. Include a picture!]
-Reply back with your first instinct…now that you know a little bit about us, can you see yourself enjoying life on <college1>’s campus?
-<Prospect Name>…what’s your favorite post-practice meal? My athletes love <mention food/snacks available on campus>.
-What are some things you’d like to know about our dorms—the place where you would be living as a student-athlete here? [OR IF ALREADY VISITED: “I’m wondering…what do you remember about our dorms when you visited? Did you like them?”]
Keep in mind that this is just for reference donot use its content in the content you will be generating.

Ensure the output aligns with the template format below:

**<College Name> <Sport>**
**{user_input}**
**TRS Messages**
In {full_months[0]} Brief description about randomly assigned topic 1 in the same format as provide above.
In {full_months[1]}: Brief description about randomly assigned topic 2 in the same format as provide above.
In {full_months[2]}: Brief description about randomly assigned topic 3 in the same format as provide above.
In {full_months[3]}: Brief description about randomly assigned topic 4 in the same format as provide above.
**For {full_months[0]}: Main Topic**
**Talking Points**
**Social Media Topic Ideas**
The email and letters should be written in a following format:
    
    **{full_months[0]}/randomly assigned topic above**
    **WEEK 1**
    **Email 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 1(Candidate Email 1) detailed content  
    <Coach Info>
    
    **{full_months[0]}/randomly assigned topic above**
    **WEEK 2**
    **Email 2**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 2(Candidate Email 2) detailed content  
    <Coach Info>

    **{full_months[0]}/randomly assigned topic above**
    **WEEK 2**
    **Parent Letter**
    
    <Prospect Name>
    Week 2 (Parent Letter) detailed content  
    <Coach Info>

    **{full_months[0]}/randomly assigned topic above**
    **WEEK 3**
    **Letter 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 3(Letter 1) content   
    <Coach Info>


    **{full_months[0]}/randomly assigned topic above**
    **WEEK 4**
    **Coach Letter**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 4(Coach Letter) detailed content   
    <Coach Info>


**For {full_months[1]}: Main Topic**
**Talking Points**
**Social Media Topic Ideas**
**Text Messaging Talking Points**
The email and letters should be written in a following format:
    
    **{full_months[1]}/randomly assigned topic above**
    **WEEK 1**
    **Email 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 1(Candidate Email 1)  detailed content  
    <Coach Info>
    
    **{full_months[1]}/randomly assigned topic above**
    **WEEK 2**
    **Parent Letter**
    
    <Prospect Name>
    Week 2 (Parent Letter) detailed content  
    <Coach Info>
    

    **{full_months[1]}/randomly assigned topic above**
    **Week 2**
    **Email 2**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 2(Candidate Email 2) detailed content  
    <Coach Info>

    **{full_months[1]}/randomly assigned topic above**
    **WEEK 3**
    **Email 3**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 3(Candidate Email 3) detailed content   
    <Coach Info>


    **{full_months[1]}/randomly assigned topic above**
    **WEEK 4**
    **Letter 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 4(Letter 1) detailed content   
    <Coach Info>

**For {full_months[2]}: Main Topic**
**Talking Points**
**Social Media Topic Ideas**
**Text Messaging Talking Points**
The email and letters should be written in a following format:
    
    **{full_months[2]}/randomly assigned topic above**
    **WEEK 1**
    **Email 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 1(Candidate Email 1) detailed content  
    <Coach Info>
    
    **{full_months[2]}/randomly assigned topic above**
    **WEEK 2**
    **Email 2**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 2(Candidate Email 2) detailed content  
    <Coach Info>

    **{full_months[2]}/randomly assigned topic above**
    **WEEK 3**
    **Parent Letter**
    
    <Prospect Name>
    Week 2 (Parent Letter) detailed content  
    <Coach Info>

    **{full_months[2]}/randomly assigned topic above**
    **WEEK 3**
    **Coach Letter**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 3(Coach Letter) detailed content   
    <Coach Info>


    **{full_months[2]}/randomly assigned topic above**
    **WEEK 4**
    **Letter 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 4(Letter 1) detailed content   
    <Coach Info>


**For {full_months[3]}: Main Topic**
**Talking Points**
**Social Media Topic Ideas**
**Text Messaging Talking Points**
The email and letters should be written in a following format:
    
    **{full_months[3]}/randomly assigned topic above**
    **WEEK 1**
    **Email 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 1(Candidate Email) detailed content  
    <Coach Info>
    
    **{full_months[3]}/randomly assigned topic above**
    **WEEK 2**
    **Parent Letter**
    
    <Prospect Name>
    Week 2 (Parent Letter) detailed content 
    <Coach Info>
    

    **{full_months[3]}/randomly assigned topic above**
    **Week 2**
    **Email 2**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 2(Candidate Email) detailed content 
    <Coach Info>

    **{full_months[3]}/randomly assigned topic above**
    **WEEK 3**
    **Email 3**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 3(Candidate Email) detailed content  
    <Coach Info>


    **{full_months[3]}/randomly assigned topic above**
    **WEEK 4**
    **Letter 1**
    **Suggested subject line**: Subject line
    <Prospect Name>
    Week 4(Letter 1) detailed content
    <Coach Info>


Make the text more conversational, and write it in a way that would make a 16 to 18 year old teenager engage with the content and be prompted to respond and interact with the coach who is sending these
Look for any grammatical errors and correct them.
VERY IMPORTANT: Place the proper spacing between paragraphs in the revised text
Keep the same layout format in place when you construct your revised text
Use the content from the {docx_content} to fill in the placeholders for the talking points, social media ideas, and text messaging points.
Ensure that the headings, subheadings, and bullet points remain organized in the final output.
Make sure every college follows the same template structure.
ONLY include the content necessary for generating the recruiting message. Do not add any extra or irrelevant details.




"""

  

    return prompt
# def prompt_generator(content):
#     prompt="""
#         Using the information and instructions provided in {docx_content}, generate a message in the following format. Make sure that the section under **TRS Messages** includes the four points listed below, and replace the placeholder values like <College> and <Sport> with the appropriate names from the {docx_content}. Ensure that **TRS Messages** along with its bullet points appear on the first page.

#         **<College> <Sport>**
#         **Sept./Oct./Nov./Dec. 2024**
#         **TRS Messages**
#         - The residence halls and general everyday life on campus for students and athletes are September’s topics, and we know from our research with your team at <College> - and others around the country - that this is a big area of interest for this generation of recruits.
#         - In October, you’ll focus on the overall athletic climate at <College>. Your messages will give your prospects a solid idea of what it’ll be like to compete for your program and be a part of the <College> campus community as an athlete and a student.
#         - In November, you’ll be focusing on the athletic facilities at <College>. We’ll combine that conversation with your training philosophy to show your recruits how you’ll get them ready to compete at the college level.
#         - In December, you’ll focus on your <Sport> team at <College>, and you’ll be doing things like exploring the team atmosphere based on the findings in your focus group survey in these messages.

#         For September, focus on residence halls and general everyday life on campus for both students and athletes. According to our research with your team at <College> and other colleges across the country, this is a key area of interest for this generation of recruits.
#         October: Highlight the overall athletic climate at <College>, showcasing what it’s like to compete and be part of the campus community as both an athlete and a student.
#         November: Emphasize the athletic facilities and training philosophy at <College>, demonstrating how recruits will be prepared for college-level competition.
#         December: Focus on the <Sport> team atmosphere at <College>, incorporating insights from the focus group survey to explore team dynamics.

#         Make sure that each month includes the following headings: **Talking Points**, **Social Media Topic Ideas**, and **Text Messaging Talking Points**. Each of these headings should have at least 5-6 bullet points.
#     """
#     return prompt

def prompt_generator_for_sonnet(content,user_input,full_months):
    prompt_temp = f"""
    Based on the provided content from {docx_content}, generate a recruiting message in the format below. Maintain the structure and include the specified headings, subheadings, and bullet points.

    **<College> <Sport>**
    **{user_input}**
    **TRS Messages**
    - The residence halls and general everyday life on campus for students and athletes are September’s topics, and we know from our research with your team at <college1> - and others around the country - that this is a big area of interest for this generation of recruits.
    - In October, you’ll focus on the overall athletic climate at <College>. Your messages will give your prospects a solid idea of what it’ll be like to compete for your program and be a part of the <College1> campus community as an athlete and a student.
    In November, you’ll be focusing on the athletic facilities at <college1>. We’ll combine that conversation with your training philosophy to show your recruits how you’ll get them ready to compete at the college level.
    - In December, you’ll focus on your <sport> team at <college1> and you’ll be doing things like exploring the team atmosphere based on the findings in your focus group survey in these messages.
    Fetch all these placeholder values from {docx_content}

    **For {full_months[0]}:**
    Highlight residence halls and daily student-athlete life. This aligns with research from <College1> and other institutions.
    **Talking Points (5-6 bullet points)**:It should often be questions or suggestions of things for the coach to send via text message to the recruit.
    **Social Media Topic Ideas (5-6 bullet points)**
    **Text Messaging Talking Points (5-6 bullet points)**:It should often be questions or suggestions of things for the coach to send via text message to the recruit.

    **For {full_months[1]}:**
    Showcase the athletic experience at <College>—what it’s like being part of the athletic community.
    **Talking Points (5-6 bullet points)**:It should often be questions or suggestions of things for the coach to send via text message to the recruit.
    **Social Media Topic Ideas (5-6 bullet points)**
    **Text Messaging Talking Points (5-6 bullet points)**:It should often be questions or suggestions of things for the coach to send via text message to the recruit.

    **For {full_months[2]}:**
    Emphasize facilities and training that prepare recruits for college-level competition.
    **Talking Points (5-6 bullet points)**: It should often be questions or suggestions of things for the coach to send via text message to the recruit.
    **Social Media Topic Ideas (5-6 bullet points)**
    **Text Messaging Talking Points (5-6 bullet points**: It should often be questions or suggestions of things for the coach to send via text message to the recruit.

    **For {full_months[3]}:**
    Focus on team dynamics and atmosphere. Include relevant insights from focus group surveys.
    **Talking Points (5-6 bullet points)**:It should often be questions or suggestions of things for the coach to send via text message to the recruit.
    **Social Media Topic Ideas (5-6 bullet points)**
    **Text Messaging Talking Points (5-6 bullet points)**:It should often be questions or suggestions of things for the coach to send via text message to the recruit.
    
    The final output should strictly follow the template format provided above.
    Ensure that the headings, subheadings, and bullet points remain organized in the final output.
    ONLY include the content necessary for generating the recruiting message. Do not add any extra or irrelevant details.
    Make the text conversational, engaging for a 16 to 18-year-old audience, prompting them to respond and interact with the coach who is sending it. Correct any grammatical errors as needed.
    VERY IMPORTANT: Ensure proper spacing between paragraphs in the revised text.

    
    


    """
    return prompt_temp

# # Function to format and save content in .docx
# def format_content(doc, content):
#     lines = content.splitlines()
#     for line in lines:
#         line = line.strip()
#         if not line:
#             continue
#         if re.match(r'^\*\*.+\*\*$', line):
#             doc.add_heading(line.replace('**', ''), level=1)
#         elif re.match(r'^\*.+\*$', line):
#             doc.add_heading(line.replace('*', ''), level=2)
#         elif line.startswith("-"):
#             if len(line) > 5:
#                 doc.add_paragraph(line[1::], style='List Bullet')
#         else:
#             doc.add_paragraph(line)


# Function to extract headings using regex from a plain string
def extract_headings(paragraphs):
    headings = []
    for para in paragraphs:
        # No need to access para.text, as para is already a string
        if re.match(r'^\*\*.*\*\*$', para.strip()):  # Check if the line is a heading
            headings.append(para.strip().strip('*'))
    return headings






def zip_folder(folder_path, output_path):
    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for root, dirs, files in os.walk(folder_path):
            for file in files:
                file_path = os.path.join(root, file)
                # Preserve folder structure inside the zip
                arcname = os.path.relpath(file_path, start=folder_path)
                zipf.write(file_path, arcname=arcname)

# Streamlit UI
st.title("Content Generator")
group = st.text_input("Enter group name:")
model_selection = st.selectbox("Choose Model", ["ChatGPT-4", "LLaMA3", "Sonnet"])
cycle = st.selectbox("Choose Cycle", ["Jan./Feb./Mar./Apr 2024", "May./Jun./July./August 2024", "Sep./Oct./Nov./Dec 2024"])
run_process = st.button("Generate Responses")

import re


user_input = cycle


# Mapping abbreviated months to full names
month_map = {
    "Jan": "January", "Feb": "February", "Mar": "March", "Apr": "April",
    "May": "May", "Jun": "June", "Jul": "July", "Aug": "August",
    "Sept": "September", "Oct": "October", "Nov": "November", "Dec": "December"
}

# Extracting abbreviated month names
abbreviations = re.findall(r'[A-Za-z]+', user_input)

# Standardizing abbreviations to match the dictionary keys (converting to title case)
full_months = [month_map.get(month.title(), month) for month in abbreviations]

# Displaying full month names
# print(full_months)

if run_process and group:
    # List files in the group folder
    response = service.files().list(q=f"name = '{group}' and mimeType = 'application/vnd.google-apps.folder'",
                                    spaces='drive').execute()
    folders = response.get('files', [])

    if not folders:
        st.error("Folder not found.")
    else:
        folder_id = folders[0]['id']
        response_folder = os.path.join(f'{group}_responses', group)  # Create Group A folder inside responses
        if not os.path.exists(response_folder):
            os.makedirs(response_folder)

        # Create a new folder 'Group A responses' in Google Drive
        responses_folder_id = create_drive_folder(service, 'Group A responses')

       
        subfolders = list_files_in_folder(folder_id)
        for folder in subfolders:
            if folder['mimeType'] == 'application/vnd.google-apps.folder':
                st.write(f"Subfolder: {folder['name']}")
                files_in_subfolder = list_files_in_folder(folder['id'])

                for file in files_in_subfolder:
                    if file['name'].endswith('.docx'):
                        st.write(f"Processing File: {file['name']}")
                        file_stream = stream_file(service, file['id'])
                        docx_content = read_docx(file_stream)

                        # Generate GPT prompt and response
                        prompt_template = prompt_generator(docx_content,user_input,full_months)
                        promptt = prompt_generator_for_sonnet(docx_content,user_input,full_months)
                        if model_selection == "ChatGPT-4":
                            response = openai.ChatCompletion.create(
                                model="gpt-4-turbo",
                                messages=[
                                    {"role": "system", "content": "You are a helpful assistant."},
                                    {"role": "user", "content": prompt_template}
                                ],
                                max_tokens=2000,
                                temperature=0.7
                            )
                            result = response.choices[0].message['content']
                            print(result)
                            # extract_headings(result)
                           # Now you can pass result.splitlines() to extract_headings since result is a string
                            all_headings = extract_headings(result.splitlines())
                            print("all_headings:", all_headings)

                            # Ensure there are enough headings to prevent IndexError
                            if len(all_headings) >= 3:
                                first_heading = all_headings[0]
                                second_heading = all_headings[1]
                                third_heading = all_headings[2]
                                headings = all_headings[3:]
                                print("first_heading:", first_heading)
                                print("second_heading:", second_heading)
                                print("third_heading:", third_heading)
                                print("headings:", headings)
                            else:
                                print("Not enough headings found in the document.")
                                first_heading = ""
                                second_heading = ""
                                third_heading = ""
                                headings = []
                            # Save response in .docx
                            college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                            doc = Document()
                            # format_content(doc, result)
                            section = doc.sections[0]
                            header = section.header

                            # Add a paragraph to the header and insert the logo image on the left
                            header_paragraph = header.paragraphs[0]
                            logo_path = 'logos_proj.jpeg'  # Replace with the actual path to your logo
                            run = header_paragraph.add_run()
                            run.add_picture(logo_path, width=Inches(4))
                            # Add the custom headings at the top of the first page
                            lines = result.splitlines()
                            if first_heading:
                                heading = doc.add_heading(first_heading, level=1)
                                heading_font = heading.runs[0].font
                                heading_font.size = Pt(24) 
                                heading_font.color.rgb = RGBColor(11,83,148)
                                heading_font.bold = True
                            if second_heading:
                                heading = doc.add_heading(second_heading, level=1)
                                heading_font = heading.runs[0].font
                                heading_font.size = Pt(24)
                                heading_font.color.rgb = RGBColor(230,145,56)
                                heading_font.bold = True
                            if third_heading:
                                heading = doc.add_heading(third_heading, level=1)
                                heading_font = heading.runs[0].font
                                heading_font.size = Pt(24)
                                heading_font.color.rgb = RGBColor(230,145,56)
                                heading_font.bold = True

                            collecting_content = True  # Start collecting content after initial headings

                            # Initialize variables
                            current_heading = None

                            # Function to check if a line is a month heading
                            def is_month_heading(text):
                                return re.match(r'^For\s\w+:', text)

                            # Function to check if the text is a heading
                            def is_heading(text):
                                return re.match(r'^\*\*.*\*\*$', text.strip())

                            # Function to get the heading text without asterisks
                            def get_heading_text(text):
                                return text.strip().strip('*')

                            # Function to check if heading is 'Talking Points'
                            def is_talking_points(heading_text):
                                return heading_text == 'Talking Points'

                            # Function to check if heading is 'Social Media Topic Ideas'
                            def is_social_media_topic_ideas(heading_text):
                                return heading_text == 'Social Media Topic Ideas'

                            # Function to check if heading is 'Text Messaging Talking Points'
                            def is_text_messaging_talking_points(heading_text):
                                return heading_text == 'Text Messaging Talking Points'

                            # Iterate over each line in the result text
                            for line in lines:
                                line_text = line.strip()
                                # print("line_text:", line_text)

                                # Check if the line is a heading
                                if is_heading(line_text):
                                    heading_text = get_heading_text(line_text)
                                    # print("heading_text:", heading_text)

                                    if heading_text in [first_heading, second_heading, third_heading]:
                                        # These headings are already added on the first page
                                        # So we skip adding them again but start collecting content if it's 'TRS Messages'
                                        collecting_content = True
                                        current_heading = heading_text
                                        # print("Skipped heading already added:", heading_text)
                                        continue

                                    # If it's a month heading
                                    elif is_month_heading(heading_text):
                                        # Always add a page break before month headings
                                        doc.add_page_break()
                                        # Add the month heading
                                        # doc.add_heading(heading_text, level=1)
                                        heading = doc.add_heading(heading_text, level=1)
                                        heading_font = heading.runs[0].font
                                        heading_font.size = Pt(16)
                                        heading_font.color.rgb = RGBColor(11,83,148)
                                        current_heading = heading_text
                                        collecting_content = True
                                        # print("Added month heading:", heading_text)

                                    elif is_talking_points(heading_text):
                                        # Add 'Talking Points' heading without page break
                                        heading = doc.add_heading(heading_text, level=1)
                                        heading_font = heading.runs[0].font
                                        heading_font.size = Pt(16)
                                        heading_font.color.rgb = RGBColor(230,145,56)
                                        current_heading = heading_text
                                        collecting_content = True
                                        # print("Added 'Talking Points' heading:", heading_text)

                                    elif is_social_media_topic_ideas(heading_text) or is_text_messaging_talking_points(heading_text):
                                        # Add a page break before these headings
                                        doc.add_page_break()
                                        doc.add_heading(heading_text, level=1)
                                        heading_font = heading.runs[0].font
                                        heading_font.size = Pt(16)
                                        heading_font.color.rgb = RGBColor(230,145,56)
                                        current_heading = heading_text
                                        collecting_content = True
                                        # print("Added heading with page break:", heading_text)

                                    else:
                                        # For any other heading, add a page break and then the heading
                                        doc.add_page_break()
                                        heading = doc.add_heading(heading_text, level=1)
                                        heading_font = heading.runs[0].font
                                        heading_font.size = Pt(16)
                                        heading_font.color.rgb = RGBColor(230,145,56)
                                        current_heading = heading_text
                                        collecting_content = True
                                        # print("Added other heading with page break:", heading_text)

                                else:
                                    # It's a normal paragraph or bullet point
                                    if collecting_content:
                                        if line.startswith("-") or re.match(r'^\d+\.', line): 
                                            if len(line) > 5: # Handle numeric or bullet points
                                                # Add bullet points
                                                paragraph = doc.add_paragraph(line[1:].strip(), style='List Bullet')
                                                if paragraph.runs:
                                                    paragraph_font = paragraph.runs[0].font
                                                    # Set font size and other properties as needed
                                                    paragraph_font.size = Pt(12)  # Example of setting the font size
                                                else:
                                                    print("Paragraph has no runs:", paragraph.text)
                                                
                                            # print("Added bullet point:", line[1:].strip())
                                        else:
                                            
                                            # Add a regular paragraph
                                            paragraph = doc.add_paragraph(line_text)
                                            if paragraph.runs:
                                                paragraph_font = paragraph.runs[0].font
                                                # Set font size and other properties as needed
                                                paragraph_font.size = Pt(12)  # Example of setting the font size
                                            else:
                                                print("Paragraph has no runs:", paragraph.text)
                                            # doc.add_paragraph(line_text)
                                            # print("Added paragraph:", line_text)
                                    else:
                                        # Skip content before any headings
                                        print("Skipped paragraph before any heading")
                            doc.save(college_response_file)
                            st.write(f"Saved: {college_response_file}")
                            uploaded_file_id = upload_to_drive(service, f"{folder['name']} response.docx", college_response_file, responses_folder_id)
                            st.write(f"Uploaded: {folder['name']} response.docx")

                        elif model_selection == "LLaMA3":
                            client = Groq(api_key=os.getenv("GROQ_API_KEY"))
                            chat_completion = client.chat.completions.create(
                                messages=[
                                    {
                                        "role": "user",
                                        "content": prompt_template,
                                    }
                                ],
                                model="llama3-8b-8192",
                            )
                            result = chat_completion.choices[0].message.content
                            print(result)
                            # Save response in .docx
                            college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                            extract_headings(result)
                            # Extract headings and subheadings from the document
                            all_headings = extract_headings(doc.paragraphs)
                            print("all_headings:", all_headings)

                            # Ensure there are enough headings to prevent IndexError
                            if len(all_headings) >= 3:
                                first_heading = all_headings[0]
                                second_heading = all_headings[1]
                                third_heading = all_headings[2]
                                headings = all_headings[3:]
                                print("first_heading:", first_heading)
                                print("second_heading:", second_heading)
                                print("third_heading:", third_heading)
                                print("headings:", headings)
                            else:
                                print("Not enough headings found in the document.")
                                first_heading = ""
                                second_heading = ""
                                third_heading = ""
                                headings = []
                            doc = Document()
                            # format_content(doc, result)
                            doc.save(college_response_file)
                            st.write(f"Saved: {college_response_file}")
                            uploaded_file_id = upload_to_drive(service, f"{folder['name']} response.docx", college_response_file, responses_folder_id)
                            st.write(f"Uploaded: {folder['name']} response.docx")

                        elif model_selection == "Sonnet":
                            client = anthropic.Anthropic(
                                api_key=os.getenv("ANTHROPIC_API_KEY"),
                            )
                            message = client.messages.create(
                                model="claude-3-5-sonnet-20240620",
                                max_tokens=2000,
                                temperature=0,
                                system="You are a helpful assistant for generating long text.",
                                messages=[
                                    {
                                        "role": "user",
                                        "content": [
                                            {
                                                "type": "text",
                                                "text": promptt
                                            }
                                        ]
                                    }
                                ]
                            )
                            result = message.content[0].text
                            print(result)
                            # Save response in .docx
                            college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                            doc = Document()
                            # format_content(doc, result)
                            doc.save(college_response_file)
                            st.write(f"Saved: {college_response_file}")
                            # Upload the generated .docx file to the 'Group A responses' folder in Google Drive
                            uploaded_file_id = upload_to_drive(service, f"{folder['name']} response.docx", college_response_file, responses_folder_id)
                            st.write(f"Uploaded: {folder['name']} response.docx")

                        # # Save response in .docx inside the Group A folder
                        # college_response_file = os.path.join(response_folder, f"{folder['name']} response.docx")
                        # doc = Document()
                        # format_content(doc, result)
                        # doc.save(college_response_file)
                        # st.write(f"Saved: {college_response_file}")

        # Zip the entire Group A folder
        zip_file_path = f"{group}_responses.zip"
        zip_folder(os.path.join(f'{group}_responses'), zip_file_path)

        # Provide download button for the ZIP file
        with open(zip_file_path, "rb") as f:
            st.download_button(
                label=f"Download {group} responses (ZIP)",
                data=f,
                file_name=f"{group}_responses.zip",
                mime="application/zip"
            )




